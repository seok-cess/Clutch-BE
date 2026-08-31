from pathlib import Path
from shutil import copy2
import zipfile
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt, RGBColor


SOURCE = Path(r"C:\Users\User\Desktop\종프 문서\종합프로젝트_기획안_3조(석세스) v3.docx")
OUTPUT = Path(r"C:\Users\User\Desktop\종프 문서\종합프로젝트_기획안_3조(석세스) v4.docx")
WORKING = Path(r"C:\Users\User\Desktop\Clutch-BE\tmp\proposal-v4\v4-working.docx")
PACKAGED = Path(r"C:\Users\User\Desktop\Clutch-BE\tmp\proposal-v4\v4-packaged.docx")
FONT = "맑은 고딕"
BLUE = RGBColor(0x1F, 0x37, 0x55)
BLACK = RGBColor(0x00, 0x00, 0x00)


def set_font(run, size, bold=False, color=BLACK):
    run.bold = bold
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color


def clear_cell(cell):
    for paragraph in list(cell.paragraphs):
        cell._element.remove(paragraph._element)


def add_paragraph(cell, before=0, after=1, line=1.05):
    paragraph = cell.add_paragraph()
    paragraph.style = "Normal"
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return paragraph


def add_labeled_paragraph(cell, label, body, label_size=10.0, body_size=9.0):
    paragraph = add_paragraph(cell, before=2.5, after=0.5)
    label_run = paragraph.add_run(f"{label}  ")
    set_font(label_run, label_size, bold=True, color=BLUE)
    body_run = paragraph.add_run(body)
    set_font(body_run, body_size, bold=False, color=BLACK)


def add_section_heading(cell, text):
    paragraph = add_paragraph(cell, before=3, after=1.5)
    run = paragraph.add_run(text)
    set_font(run, 10.5, bold=True, color=BLUE)


def add_body(cell, text):
    paragraph = add_paragraph(cell, before=0, after=1)
    run = paragraph.add_run(text)
    set_font(run, 9.5, bold=False, color=BLACK)


def add_hyperlink(paragraph, text, url):
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    fonts.set(qn("w:eastAsia"), FONT)
    run_properties.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "16")
    run_properties.append(size)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)
    run.append(run_properties)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_reference_paragraph(cell):
    paragraph = add_paragraph(cell, before=0, after=1, line=1.0)
    label = paragraph.add_run("참고 자료  ")
    set_font(label, 8.0, bold=True, color=BLUE)
    add_hyperlink(
        paragraph,
        "LCK, 역대 최고 시청 지표 달성",
        "https://v.daum.net/v/xNJM3Sd8WI",
    )
    separator = paragraph.add_run(" · ")
    set_font(separator, 8.0, bold=False, color=BLACK)
    add_hyperlink(
        paragraph,
        "LCK, 글로벌 콘텐츠로 성장",
        "https://www.hankyung.com/article/202410082353i",
    )


def replace_paragraph(paragraph, text):
    first = paragraph.runs[0]
    first.text = text
    for run in paragraph.runs[1:]:
        run._element.getparent().remove(run._element)


copy2(SOURCE, WORKING)
document = Document(WORKING)
table = document.tables[0]

# 프로젝트 목적: 대상·가치, 핵심 과제, 핵심/확장 범위, 현실적 구성 제약을 명시한다.
purpose_cell = table.rows[3].cells[1]
clear_cell(purpose_cell)
add_labeled_paragraph(
    purpose_cell,
    "도메인 선정 배경",
    "LCK는 평균 분당 시청자가 44.9만~63.4만 명에 이를 만큼 다수의 이용자가 동일 경기를 실시간으로 시청하는 도메인이다. 첫 킬·펜타킬과 같은 순간 이벤트에는 이용자 반응이 집중될 수 있으며, 평균 시청자의 일부만 참여해도 2만 건 규모에 도달한다. 따라서 e스포츠는 재고 1만 장에 최대 2만 VU가 몰리는 선착순 쿠폰 발급 시나리오를 적용하기에 적합하다.",
    label_size=10.5,
    body_size=9.5,
)
add_reference_paragraph(purpose_cell)
add_labeled_paragraph(
    purpose_cell,
    "대상·가치",
    "경기 시청자에게는 실시간 참여 보상과 이벤트 경험을, 운영자에게는 대규모 선착순 쿠폰을 안전하게 운영·검증할 수 있는 수단을 제공한다.",
    label_size=10.5,
    body_size=9.5,
)
add_labeled_paragraph(
    purpose_cell,
    "핵심 과제",
    "재고 10,000장에 60초 ramp-up으로 최대 20,000 VU가 요청하는 조건에서 동일 회차 중복·초과 발급 없이 MySQL에 실제 쿠폰을 확정한다.",
    label_size=10.5,
    body_size=9.5,
)
add_labeled_paragraph(
    purpose_cell,
    "기능 범위",
    "선착순 발급, 개인정보 마스킹, 100만 사용자·300만 발급 이력 전체 정합성과 재실행 결정론 검증을 핵심으로 하고, 경기 트리거·SSE·시청 포인트·세트 배팅은 확장 기능으로 구성한다.",
    label_size=10.5,
    body_size=9.5,
)
add_labeled_paragraph(
    purpose_cell,
    "구현 범위",
    "교육용 가상 데이터와 Replay STUB·Mock 연동, 단일 CLUTCH API·MySQL·Redis·Kafka 구성을 기준으로 비용과 범위를 통제하고 운영 통계와 검증 이력을 제공한다.",
    label_size=10.5,
    body_size=9.5,
)

# 프로젝트 수행 방향: 실제 수행 내용과 필수 검증·문서화가 일정에서 보이도록 보정한다.
schedule_cell = table.rows[4].cells[1]
replace_paragraph(
    schedule_cell.paragraphs[5],
    "8/24~8/27  동기 발급 확정, 성공 수량 집계 분리, Redis 장애 복구, 100만/300만 전체 정합성 검증",
)
replace_paragraph(
    schedule_cell.paragraphs[6],
    "8/28~8/31  Replay·첫 킬/펜타킬 트리거, 관리자 통계, 20,000 VU 테스트, 실행 가이드·최종 문서화",
)
replace_paragraph(
    schedule_cell.paragraphs[8],
    "석종수(조장): 프로젝트 관리, CI/CD·인프라, 부하 테스트·모니터링, 전체 정합성 검증·문서화",
)

# 필수 기능: 과제 핵심 기능과 서비스 확장 기능을 명시적으로 분리한다.
feature_cell = table.rows[6].cells[1]
clear_cell(feature_cell)
add_section_heading(feature_cell, "과제 핵심 기능")
add_labeled_paragraph(
    feature_cell,
    "선착순 쿠폰 발급",
    "Redis Lua로 재고 차감과 회차별 중복을 원자 처리하고, MySQL에서 실제 사용자 쿠폰과 결과 Outbox를 한 트랜잭션으로 확정한다.",
)
add_labeled_paragraph(
    feature_cell,
    "대량 데이터·개인정보",
    "가상 사용자 100만 명과 발급 요청 이력 300만 건 이상을 적재하고, 관리자 응답의 이름·이메일·전화번호를 마스킹한다.",
)
add_labeled_paragraph(
    feature_cell,
    "상태·전체 정합성",
    "발급·사용·취소·만료와 반복/동시 상태 변경을 멱등 처리하고, 300만 건 전체의 참조·중복·재고·상태 불일치를 검증한다.",
)
add_labeled_paragraph(
    feature_cell,
    "결정론·재현",
    "동일 스냅샷의 건수·ID 범위·fingerprint로 재실행 결과를 비교하고, k6 실행 가이드와 독립 SQL·관리자 비동기 검증 API를 제공한다.",
)
add_section_heading(feature_cell, "서비스 확장 기능")
add_labeled_paragraph(
    feature_cell,
    "경기·이벤트",
    "실제 API와 Replay STUB을 전환해 경기·세트 결과와 순위를 제공하고 첫 킬·펜타킬을 감지해 쿠폰 회차를 연다.",
)
add_labeled_paragraph(
    feature_cell,
    "사용자 참여",
    "회차별 신청 가능 시간과 실시간 잔여 재고(SSE), 내 쿠폰 상태를 제공하며 5분 시청 보상과 세트 승패 배팅을 연계한다.",
)
add_labeled_paragraph(
    feature_cell,
    "운영·외부 연동",
    "관리자 이벤트 제어와 성공·거절·Kafka 오류 통계를 제공하고, 알림 등 외부 연동은 Mock으로 대체한다.",
)

# 포함 기술: 기획안의 시제에 맞춰 기술 적용 방향, 예상 효과와 검증 계획을 제시한다.
tech_cell = table.rows[7].cells[1]
clear_cell(tech_cell)
add_section_heading(tech_cell, "정합성과 장애 대응")
add_body(tech_cell, "Redis Lua로 재고·중복을 선판단하고 Redis 불확실 상태에서는 fail-closed로 발급을 중단한다. MySQL의 실제 user_coupon을 최종 기준으로 Redis 재고·당첨자·발급 컨텍스트를 복구한다.")
add_body(tech_cell, "성공 수량 갱신을 핵심 발급 트랜잭션에서 분리하고 MySQL named lock 기반 비동기 집계를 적용해 공통 성공 수량 행의 경합을 완화한다.")
add_section_heading(tech_cell, "메시징과 운영 관측")
add_body(tech_cell, "MySQL Outbox와 Kafka로 확정 발급 결과를 후속 처리하며 messageId·원본 Kafka 좌표로 중복 집계를 방지하고 재시도·DLT를 통계에 반영한다.")
add_body(tech_cell, "Actuator/Micrometer, Prometheus, Grafana와 MySQL/Redis Exporter로 애플리케이션·저장소·k6 지표를 함께 관찰한다.")
add_section_heading(tech_cell, "예상 효과")
add_body(tech_cell, "재고 10,000장에 60초 ramp-up으로 최대 20,000 VU가 신청해도 초과·중복 발급 없이 10,000건을 확정하고 나머지는 정상 품절로 처리할 것으로 예상한다. 신청 API p95는 5초 미만, 전송 실패는 0건을 목표로 한다.")
add_body(tech_cell, "가상 사용자 100만 명과 발급 요청 이력 300만 건 전체에서 참조·상태·중복·재고 불일치가 0건이고, 동일 스냅샷 재실행 시 건수·ID 범위·fingerprint가 일치할 것으로 기대한다.")
add_section_heading(tech_cell, "검증 계획")
add_body(tech_cell, "k6로 20,000 VU 부하를 재현하고 관리자 API와 MySQL 실제 쿠폰 수를 교차 확인한다. 독립 SQL과 관리자 비동기 검증 API로 300만 건 전체를 검사하며 Outbox/Kafka backlog와 동일 사용자의 동시 중복 신청도 별도 시나리오로 확인한다.")
add_section_heading(tech_cell, "예상 한계와 대응")
add_body(tech_cell, "Docker NAT 등 부하 발생기 경로가 병목이 될 수 있으므로 네이티브 실행 결과와 비교한다. 기본 단일 애플리케이션·MySQL·Redis·Kafka 구성을 범위로 두고, Redis 장애 시에는 발급을 차단한 뒤 MySQL 기준으로 복구한다.")

document.save(WORKING)

# python-docx가 수정하지 않은 styles/footer/settings까지 재직렬화하지 않도록,
# 원본 패키지의 word/document.xml만 작업본의 본문 XML로 교체한다.
with zipfile.ZipFile(WORKING, "r") as working_zip:
    edited_parts = {
        "word/document.xml": working_zip.read("word/document.xml"),
        "word/_rels/document.xml.rels": working_zip.read("word/_rels/document.xml.rels"),
    }

with zipfile.ZipFile(SOURCE, "r") as source_zip, zipfile.ZipFile(PACKAGED, "w") as output_zip:
    for entry in source_zip.infolist():
        payload = edited_parts.get(entry.filename, source_zip.read(entry.filename))
        output_zip.writestr(entry, payload)

PACKAGED.replace(OUTPUT)
print(OUTPUT)
